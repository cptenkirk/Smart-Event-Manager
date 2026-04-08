from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from docx import Document as DocxDocument
import os
from langchain_text_splitters import RecursiveCharacterTextSplitter # NEU

def get_full_docx_text(filename):
    doc = DocxDocument(filename)
    return '\n'.join([para.text for para in doc.paragraphs])

# --- 1. Daten laden ---
docx_file_path = r"./data/INTRO.docx" # Pfad angepasst für Docker/Linux
full_text_content = get_full_docx_text(docx_file_path)

# --- 2. TEXT SPLITTING (Das fehlende Puzzleteil) ---
# Wir teilen den Text in 500-Zeichen-Blöcke mit Überlappung
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
chunks = text_splitter.split_text(full_text_content)

# Chunks in Document-Objekte umwandeln
documents = [Document(page_content=t, metadata={"source": docx_file_path}) for t in chunks]

# --- 3. Initialisierung ---
# WICHTIG: Nutze nomic-embed-text für Geschwindigkeit!
embeddings = OllamaEmbeddings(model="nomic-embed-text")
db_location = "chroma_db"

add_documents = not os.path.exists(db_location) 

vector_store = Chroma(
    collection_name="lab_collection",
    persist_directory=db_location,
    embedding_function=embeddings
)

if add_documents: 
    print(f"Erstelle Vektorspeicher und füge {len(documents)} Chunks hinzu...")
    vector_store.add_documents(documents=documents)
else:
    print(f"Lade bestehenden Vektorspeicher...")

retriever = vector_store.as_retriever(search_kwargs={"k": 3})
